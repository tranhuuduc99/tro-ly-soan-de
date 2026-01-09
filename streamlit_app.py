import streamlit as st
import requests
import json
import time
from docx import Document
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import random

# --- CẤU HÌNH TRANG ---
st.set_page_config(page_title="Phần mềm Soạn đề 2.0", page_icon="🛠️", layout="wide")

# --- DANH SÁCH KEY CỦA BẠN ---
API_KEYS = [
    "AIzaSyBY29kMfQWCB7ASsBrWcPHKn8EG8kYq_Bc",  # Key 1
]

# --- 1. HÀM TỰ ĐỘNG TÌM ĐÚNG TÊN MODEL (Fix 404) ---
def get_working_model(api_key):
    """
    Hàm này sẽ hỏi Google xem Key này dùng được model nào.
    Nó sẽ ưu tiên Flash -> Pro -> 1.0 -> 1.5 để tránh lỗi 404.
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
    try:
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            data = response.json()
            models = [m['name'].replace('models/', '') for m in data.get('models', []) if 'generateContent' in m.get('supportedGenerationMethods', [])]
            
            # Ưu tiên theo thứ tự ngon - bổ - rẻ
            if 'gemini-1.5-flash' in models: return 'gemini-1.5-flash'
            if 'gemini-1.5-pro' in models: return 'gemini-1.5-pro'
            if 'gemini-1.0-pro' in models: return 'gemini-1.0-pro'
            if 'gemini-pro' in models: return 'gemini-pro'
            
            # Nếu không tìm thấy cái ưu tiên, lấy cái đầu tiên tìm được
            if models: return models[0]
            
    except:
        pass
    # Fallback cuối cùng
    return "gemini-pro"

# --- 2. HÀM GỌI API ĐA LUỒNG ---
def call_gemini_auto(prompt):
    valid_keys = API_KEYS.copy()
    random.shuffle(valid_keys)
    
    errors = []

    for i, key in enumerate(valid_keys):
        clean_key = key.strip()
        
        # BƯỚC QUAN TRỌNG: Tìm model đúng cho Key này trước khi gọi
        correct_model = get_working_model(clean_key)
        
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{correct_model}:generateContent?key={clean_key}"
        headers = {'Content-Type': 'application/json'}
        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "safetySettings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
            ]
        }

        try:
            # st.toast(f"Đang thử Key {i+1} với model {correct_model}...", icon="🤖")
            response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=30)
            
            if response.status_code == 200:
                return response.json()['candidates'][0]['content']['parts'][0]['text']
            elif response.status_code == 429:
                errors.append(f"Key {i+1} quá tải")
                continue 
            elif response.status_code == 404:
                errors.append(f"Key {i+1} lỗi 404 (Sai model)")
                continue
            else:
                errors.append(f"Key {i+1} lỗi {response.status_code}")
                continue

        except Exception as e:
            errors.append(f"Key {i+1} lỗi mạng")
            continue

    return f"❌ KHÔNG THỂ TẠO ĐỀ. Chi tiết lỗi:\n{'; '.join(errors)}"

# --- 3. HÀM XỬ LÝ FILE ---
def read_file(uploaded_file):
    try:
        if uploaded_file.name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(uploaded_file)
            text = "".join([page.extract_text() or "" for page in reader.pages])
            return text
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        return ""
    except: return "Lỗi đọc file."

# --- 4. HÀM TẠO FILE WORD ---
def create_word(content, topic, grade_info):
    doc = Document()
    h = doc.add_heading('TRƯỜNG PTDTBT THCS MÙN CHUNG', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f'ĐỀ KIỂM TRA: {topic.upper()}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Lớp: {grade_info}      Họ và tên: .................................................")
    doc.add_paragraph("-" * 60)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 5. GIAO DIỆN CHÍNH ---
st.markdown("<h1 style='text-align: center; color: #004d99;'>🛠️ PHẦN MỀM SOẠN ĐỀ 2.0</h1>", unsafe_allow_html=True)

with st.sidebar:
    st.success(f"✅ Đã nạp {len(API_KEYS)} Key")
    st.info("Đã bật chế độ tự động tìm đúng tên Model để tránh lỗi 404.")

tab1, tab2, tab3 = st.tabs(["⚡ Soạn Theo Chủ Đề", "📂 Soạn Từ File", "📥 Tải Về"])

# --- TAB 1 ---
with tab1:
    col1, col2 = st.columns(2)
    with col1:
        topic = st.text_input("Chủ đề:", "Vợ chồng A Phủ")
        grade = st.selectbox("Khối:", [str(i) for i in range(1, 13)], index=11) 
        subject = st.text_input("Môn:", "Ngữ Văn")
        level = st.select_slider("Độ khó:", ["Cơ bản", "Khá", "Nâng cao", "HSG"])
    with col2:
        st.write("--- Cấu trúc ---")
        num_mc = st.number_input("TN (4 đáp án):", 0, 50, 6)
        num_tf = st.number_input("TN Đúng/Sai:", 0, 20, 2)
        num_tl = st.number_input("Tự luận:", 0, 10, 1)

    if st.button("🚀 SOẠN ĐỀ NGAY", type="primary", use_container_width=True):
        reqs = []
        if num_mc > 0: reqs.append(f"- {num_mc} câu Trắc nghiệm (4 đáp án).")
        if num_tf > 0: reqs.append(f"- {num_tf} câu Đúng/Sai (dạng chùm).")
        if num_tl > 0: reqs.append(f"- {num_tl} câu Tự luận.")
        
        prompt = (
            f"Bạn là giáo viên môn {subject} lớp {grade}. Chủ đề: '{topic}'. Độ khó: {level}.\n"
            f"Yêu cầu cấu trúc:\n" + "\n".join(reqs) + 
            f"\nCung cấp Đề bài và Đáp án chi tiết tách biệt."
        )
        
        with st.spinner("Đang tìm model phù hợp và soạn đề..."):
            res = call_gemini_auto(prompt)
            st.session_state['res'] = res
            st.session_state['top'] = topic
            st.session_state['gr'] = grade
            if "❌" in res: st.error(res)
            else: st.success("Thành công!"); st.write(res)

# --- TAB 2 ---
with tab2:
    f = st.file_uploader("Tải tài liệu:", type=['pdf','docx'])
    grade_file = st.selectbox("Lớp:", [str(i) for i in range(1, 13)], index=8, key='gr_file')
    c1, c2, c3 = st.columns(3)
    n_mc_f = c1.number_input("SL TN:", 0, 50, 10)
    n_tf_f = c2.number_input("SL Đ/S:", 0, 20, 2)
    n_tl_f = c3.number_input("SL TL:", 0, 10, 1)

    if st.button("🚀 PHÂN TÍCH", use_container_width=True):
        if f:
            with st.spinner("Đang xử lý..."):
                content = read_file(f)
                reqs = []
                if n_mc_f > 0: reqs.append(f"- {n_mc_f} câu Trắc nghiệm.")
                if n_tf_f > 0: reqs.append(f"- {n_tf_f} câu Đúng/Sai.")
                if n_tl_f > 0: reqs.append(f"- {n_tl_f} câu Tự luận.")
                
                prompt = f"Dựa vào văn bản: '{content[:15000]}'... Soạn đề lớp {grade_file}:\n" + "\n".join(reqs) + "\nCó đáp án chi tiết."
                res = call_gemini_auto(prompt)
                st.session_state['res'] = res
                st.session_state['top'] = f.name
                st.session_state['gr'] = grade_file
                if "❌" in res: st.error(res)
                else: st.success("Xong!"); st.write(res)

# --- TAB 3 ---
with tab3:
    if 'res' in st.session_state:
        txt = st.text_area("Nội dung:", st.session_state['res'], height=300)
        docx = create_word(txt, st.session_state['top'], st.session_state.get('gr', ''))
        st.download_button("📥 TẢI WORD", docx, f"{st.session_state['top']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
